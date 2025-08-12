# parse_doc.py (DEBUG version)
"""
Debug version: Extract text (and images) from a .docx into a structured JSON,
with verbose printing so we can see why parsed.json is not being created.
"""
import argparse
import zipfile
import os
import re
import traceback
from docx import Document

# make imports robust regardless of how script is run
import sys
SCRIPT_DIR = os.path.dirname(__file__)
if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)
# fallback: ensure project root is on path too
PROJ_ROOT = os.path.dirname(SCRIPT_DIR)
if PROJ_ROOT not in sys.path:
    sys.path.insert(0, PROJ_ROOT)

# Try multiple import styles
try:
    from utils import save_json, ensure_dir
except Exception:
    try:
        from src.utils import save_json, ensure_dir
    except Exception as e:
        print("ERROR importing utils:", e)
        traceback.print_exc()
        raise

def extract_images_from_docx(docx_path, out_dir):
    ensure_dir(out_dir)
    imgs = []
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            names = z.namelist()
            for item in names:
                if item.startswith("word/media/"):
                    z.extract(item, out_dir)
        media_dir = os.path.join(out_dir, "word", "media")
        if os.path.isdir(media_dir):
            for fname in sorted(os.listdir(media_dir)):
                src = os.path.join(media_dir, fname)
                dst = os.path.join(out_dir, fname)
                try:
                    os.replace(src, dst)
                except Exception:
                    try:
                        os.rename(src, dst)
                    except Exception:
                        pass
                imgs.append(dst)
            # attempt to remove empty directories
            try:
                os.removedirs(media_dir)
            except Exception:
                pass
    except Exception as e:
        print("Warning: could not extract images:", e)
        traceback.print_exc()
    return imgs

def extract_text_questions(docx_path):
    doc = Document(docx_path)
    full_text = []
    para_count = 0
    for p in doc.paragraphs:
        para_count += 1
        # preserve paragraph text
        full_text.append(p.text)
    joined = "\n".join(full_text)

    print(f"Paragraphs found: {para_count}")
    if para_count <= 20:
        print("Paragraphs (raw):")
        for i, p in enumerate(doc.paragraphs, 1):
            print(f"  [{i}] {repr(p.text)}")

    # split on lines that look like numbered questions (e.g. "1. " or "1) ")
    splits = re.split(r'\n(?=\s*\d+[\.\)]\s+)', joined)
    questions = [s.strip() for s in splits if s.strip()]

    print("Joined text (first 400 chars):")
    print(joined[:400].replace('\n','\\n'))
    print(f"Detected {len(questions)} question segments.")
    if len(questions) > 0:
        for i,q in enumerate(questions,1):
            print(f"--- Q{i} (first 200 chars) ---")
            print(q[:200].replace('\n','\\n'))
    return questions, joined

def main():
    try:
        parser = argparse.ArgumentParser()
        parser.add_argument("--input", required=True)
        parser.add_argument("--out", required=True)
        args = parser.parse_args()

        print("Running parse_doc.py (debug)")
        print("CWD:", os.getcwd())
        print("Script dir:", SCRIPT_DIR)
        print("Project root:", PROJ_ROOT)
        print("Input arg:", args.input)
        print("Input abs:", os.path.abspath(args.input))
        print("Out arg:", args.out)
        print("Out abs:", os.path.abspath(args.out))
        print("Input exists?", os.path.exists(args.input))
        print("Have permissions to write to output folder? (attempting to create)")

        out_dir = os.path.dirname(args.out) or "."
        ensure_dir(out_dir)
        media_out = os.path.join(out_dir, "media")
        ensure_dir(media_out)

        # abort early if input file not present
        if not os.path.exists(args.input):
            print("ERROR: Input file does not exist. Exiting.")
            return

        imgs = extract_images_from_docx(args.input, out_dir)
        questions, raw_text = extract_text_questions(args.input)
        data = {
            "source_file": os.path.abspath(args.input),
            "raw_text": raw_text,
            "questions": questions,
            "extracted_images": imgs
        }
        # print lengths for debug
        print("Saving JSON... (lengths) raw_text:", len(raw_text), "questions:", len(questions), "images:", len(imgs))
        save_json(data, args.out)
        print("Parsed. Saved to:", os.path.abspath(args.out))
        print("Found images:", imgs)
    except Exception as e:
        print("Unhandled exception in parse_doc.py:")
        traceback.print_exc()

if __name__ == "__main__":
    main()
