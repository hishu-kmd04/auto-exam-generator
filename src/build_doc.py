# build_doc.py
"""
Assemble the final Word doc in the required 'Question Output Format'.
Usage:
  python build_doc.py --input output/questions.json --images output/images/ --out output/result.docx
"""
import sys, os
sys.path.append(os.path.dirname(os.path.dirname(__file__)))
import argparse
from docx import Document
from docx.shared import Inches
from utils import load_json, ensure_dir


def insert_question_block(doc, q, image_path=None):
    # Title (as @title)
    doc.add_paragraph(f"@title {q.get('title','')}")
    doc.add_paragraph(f"@description {q.get('description','')}")
    doc.add_paragraph("")  # blank
    # MCQ block
    doc.add_paragraph("// Use this block for each question when adding Multiple Choice Questions (MCQ)")
    doc.add_paragraph(f"@question {q.get('question','')}")
    doc.add_paragraph(f"@instruction {q.get('instruction','')}")
    doc.add_paragraph(f"@difficulty {q.get('difficulty','')}")
    doc.add_paragraph(f"@Order {q.get('order','')}")
    for opt in q.get("options", []):
        doc.add_paragraph(f"@option {opt}")
    doc.add_paragraph(f"@@option {q.get('correct_answer','')}")
    doc.add_paragraph(f"@option {q.get('options')[-1] if q.get('options') else ''}")
    doc.add_paragraph(f"@explanation")
    doc.add_paragraph(q.get("explanation",""))
    doc.add_paragraph(f"@subject {q.get('subject','')}")
    doc.add_paragraph(f"@unit {q.get('unit','')}")
    doc.add_paragraph(f"@topic {q.get('topic','')}")
    doc.add_paragraph(f"@plusmarks {q.get('plusmarks','1')}")
    if image_path and os.path.exists(image_path):
        # add a caption and insert
        doc.add_paragraph("Figure:")
        # insert scaled image
        try:
            doc.add_picture(image_path, width=Inches(3.5))
        except Exception:
            doc.add_paragraph(f"[Image could not be inserted: {image_path}]")
    # page break between questions
    doc.add_page_break()

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--images", required=False, default=None)
    parser.add_argument("--out", required=True)
    args = parser.parse_args()

    data = load_json(args.input)
    images_map = {}
    if args.images:
        # map order -> image path by looking through folder
        for fname in os.listdir(args.images):
            # expecting filenames like '..._table.png' or '..._balls.png'
            # heuristics: parse trailing order if present, otherwise use name index
            images_map[fname] = os.path.join(args.images, fname)
    doc = Document()
    doc.add_heading("Auto-generated Questions", level=1)
    for q in data.get("questions", []):
        img_path = None
        # choose image by searching images_map keys for q title
        if args.images:
            # naive match
            for fname, path in images_map.items():
                if q.get('title','').replace(" ", "_")[:20].lower() in fname.lower():
                    img_path = path
                    break
            # fallback match by order
            if not img_path:
                # find file with q order number
                for fname, path in images_map.items():
                    if str(q.get('order')) in fname:
                        img_path = path
                        break
        insert_question_block(doc, q, image_path=img_path)
    ensure_dir(os.path.dirname(args.out) or ".")
    doc.save(args.out)
    print("Saved final doc to", args.out)

if __name__ == "__main__":
    main()
